from fastapi import FastAPI, Depends, Request, Form, HTTPException, status, UploadFile, File
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse, StreamingResponse
from sqlalchemy.orm import Session
import hashlib
from datetime import date, datetime
import os, io, shutil

from database import engine, Base, get_db
import models

# Excel / Word export
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor

# Create DB tables
Base.metadata.create_all(bind=engine)

# Ensure upload dirs exist
UPLOAD_DIR = "static/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

app = FastAPI(title="ИС учета проектов автоматизации")
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# ─── Helpers ─────────────────────────────────────────────────────────────────

def get_password_hash(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(plain: str, hashed: str) -> bool:
    return get_password_hash(plain) == hashed

def get_current_user(request: Request, db: Session):
    uid = request.cookies.get("user_id")
    if uid:
        return db.query(models.User).filter(models.User.id == int(uid)).first()
    return None

def notify(db: Session, user_id: int, message: str, link: str = None):
    n = models.Notification(user_id=user_id, message=message, link=link)
    db.add(n)

# ─── Startup seed ─────────────────────────────────────────────────────────────

@app.on_event("startup")
def create_initial_data():
    db = next(get_db())

    if db.query(models.Role).count() == 0:
        db.add_all([models.Role(name=n) for n in ("head","pm","employee","customer")])
        db.commit()

    if db.query(models.ProjectStatus).count() == 0:
        db.add_all([models.ProjectStatus(name=n) for n in ("Новый","В работе","Завершен","Приостановлен")])
        db.commit()

    if db.query(models.TaskStatus).count() == 0:
        db.add_all([models.TaskStatus(name=n) for n in ("К выполнению","В процессе","Выполнено","Отменено")])
        db.commit()

    if db.query(models.DocumentType).count() == 0:
        db.add_all([models.DocumentType(name=n) for n in ("ТЗ","Отчет","Акт","Договор","Правки")])
        db.commit()

    if db.query(models.User).count() == 0:
        rh = db.query(models.Role).filter_by(name="head").first().id
        rp = db.query(models.Role).filter_by(name="pm").first().id
        re = db.query(models.Role).filter_by(name="employee").first().id
        rc = db.query(models.Role).filter_by(name="customer").first().id
        db.add_all([
            models.User(username="head", password_hash=get_password_hash("123"), full_name="Иванов И.И. (Руководитель)", role_id=rh, position="Директор", email="head@company.ru"),
            models.User(username="pm",   password_hash=get_password_hash("123"), full_name="Петров П.П. (РП)",           role_id=rp, position="Менеджер проекта", email="pm@company.ru"),
            models.User(username="emp",  password_hash=get_password_hash("123"), full_name="Сидоров С.С. (Исполнитель)", role_id=re, position="Разработчик", email="emp@company.ru"),
            models.User(username="cust", password_hash=get_password_hash("123"), full_name="ООО Вектор (Заказчик)",      role_id=rc, position="Представитель заказчика", email="cust@vector.ru"),
        ])
        db.commit()

# ─── Auth ─────────────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return RedirectResponse(url="/projects")

@app.get("/login", response_class=HTMLResponse)
async def login_get(request: Request):
    return templates.TemplateResponse("login.html", {"request": request})

@app.post("/login")
async def login_post(request: Request, username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.username == username).first()
    if not user or not verify_password(password, user.password_hash):
        return templates.TemplateResponse("login.html", {"request": request, "error": "Неверный логин или пароль"})
    resp = RedirectResponse(url="/projects", status_code=status.HTTP_302_FOUND)
    resp.set_cookie("user_id", str(user.id))
    resp.set_cookie("user_role", user.role.name)
    return resp

@app.get("/logout")
async def logout():
    resp = RedirectResponse(url="/login")
    resp.delete_cookie("user_id")
    resp.delete_cookie("user_role")
    return resp

# ─── Profile ──────────────────────────────────────────────────────────────────

@app.get("/profile", response_class=HTMLResponse)
async def profile_get(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user: return RedirectResponse(url="/login")
    return templates.TemplateResponse("profile.html", {"request": request, "user": user})

@app.post("/profile")
async def profile_post(
    request: Request,
    full_name: str = Form(...),
    email: str = Form(""),
    phone: str = Form(""),
    position: str = Form(""),
    db: Session = Depends(get_db)
):
    user = get_current_user(request, db)
    if not user: return RedirectResponse(url="/login")
    user.full_name = full_name
    user.email = email
    user.phone = phone
    user.position = position
    db.commit()
    return RedirectResponse(url="/profile?saved=1", status_code=status.HTTP_302_FOUND)

# ─── Projects ─────────────────────────────────────────────────────────────────

@app.get("/projects", response_class=HTMLResponse)
async def list_projects(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user: return RedirectResponse(url="/login")
    rn = user.role.name
    if rn in ("head","pm"):
        projects = db.query(models.Project).all()
    elif rn == "customer":
        projects = db.query(models.Project).filter(models.Project.customer_id == user.id).all()
    else:
        projects = db.query(models.Project).join(models.Task).filter(models.Task.assignee_id == user.id).distinct().all()
    return templates.TemplateResponse("projects.html", {"request": request, "user": user, "projects": projects})

@app.get("/projects/new", response_class=HTMLResponse)
async def new_project_get(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user or user.role.name not in ("head","pm"): return RedirectResponse(url="/projects")
    customers = db.query(models.User).join(models.Role).filter(models.Role.name == "customer").all()
    pms = db.query(models.User).join(models.Role).filter(models.Role.name.in_(("pm","head"))).all()
    return templates.TemplateResponse("project_form.html", {"request": request, "user": user, "customers": customers, "pms": pms})

@app.post("/projects/new")
async def new_project_post(
    request: Request,
    title: str = Form(...),
    description: str = Form(...),
    customer_id: int = Form(...),
    pm_id: int = Form(None),
    start_date: date = Form(...),
    end_date: date = Form(None),
    planned_budget: float = Form(0.0),
    db: Session = Depends(get_db)
):
    user = get_current_user(request, db)
    if not user or user.role.name not in ("head","pm"): return RedirectResponse(url="/login")
    s = db.query(models.ProjectStatus).filter_by(name="Новый").first()
    proj = models.Project(
        title=title, description=description, customer_id=customer_id,
        pm_id=pm_id or user.id, start_date=start_date, end_date=end_date,
        planned_budget=planned_budget, status_id=s.id
    )
    db.add(proj)
    db.commit()
    db.refresh(proj)
    notify(db, customer_id, f"Создан проект «{title}». Вы можете отслеживать его в системе.", f"/projects/{proj.id}")
    db.commit()
    return RedirectResponse(url="/projects", status_code=status.HTTP_302_FOUND)

@app.get("/projects/{project_id}", response_class=HTMLResponse)
async def project_detail(request: Request, project_id: int, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user: return RedirectResponse(url="/login")
    proj = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not proj: raise HTTPException(404, "Проект не найден")
    if user.role.name == "customer" and proj.customer_id != user.id:
        return RedirectResponse(url="/projects")
    # Mark notifications read for this project
    db.query(models.Notification).filter(
        models.Notification.user_id == user.id,
        models.Notification.link == f"/projects/{project_id}",
        models.Notification.is_read == False
    ).update({"is_read": True})
    db.commit()
    return templates.TemplateResponse("project_detail.html", {"request": request, "user": user, "project": proj})

# Project chat / messages
@app.post("/projects/{project_id}/message")
async def post_project_message(
    request: Request,
    project_id: int,
    content: str = Form(...),
    is_revision: bool = Form(False),
    db: Session = Depends(get_db)
):
    user = get_current_user(request, db)
    if not user: return RedirectResponse(url="/login")
    proj = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not proj: raise HTTPException(404)
    msg = models.ProjectMessage(project_id=project_id, sender_id=user.id, content=content, is_revision=is_revision)
    db.add(msg)
    # Notify the other party
    if user.role.name == "customer":
        if proj.pm_id:
            notify(db, proj.pm_id, f"Новое сообщение от заказчика в проекте «{proj.title}»", f"/projects/{project_id}")
    elif proj.customer_id:
        notify(db, proj.customer_id, f"Новое сообщение по проекту «{proj.title}»", f"/projects/{project_id}")
    db.commit()
    return RedirectResponse(url=f"/projects/{project_id}", status_code=status.HTTP_302_FOUND)

# Upload document to project
@app.post("/projects/{project_id}/upload")
async def upload_document(
    request: Request,
    project_id: int,
    doc_title: str = Form(...),
    doc_type_id: int = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    user = get_current_user(request, db)
    if not user: return RedirectResponse(url="/login")
    proj = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not proj: raise HTTPException(404)
    # Save file
    ext = os.path.splitext(file.filename)[1]
    fname = f"proj_{project_id}_{int(datetime.utcnow().timestamp())}{ext}"
    fpath = os.path.join(UPLOAD_DIR, fname)
    with open(fpath, "wb") as f:
        shutil.copyfileobj(file.file, f)
    doc = models.Document(project_id=project_id, title=doc_title, document_type_id=doc_type_id,
                          file_path=fname, author_id=user.id)
    db.add(doc)
    db.commit()
    return RedirectResponse(url=f"/projects/{project_id}", status_code=status.HTTP_302_FOUND)

# Download document
@app.get("/documents/{doc_id}/download")
async def download_doc(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(models.Document).filter(models.Document.id == doc_id).first()
    if not doc or not doc.file_path: raise HTTPException(404)
    full = os.path.join(UPLOAD_DIR, doc.file_path)
    if not os.path.exists(full): raise HTTPException(404)
    return FileResponse(full, filename=doc.file_path)

# ─── Tasks ────────────────────────────────────────────────────────────────────

@app.get("/tasks", response_class=HTMLResponse)
async def list_tasks(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user: return RedirectResponse(url="/login")
    rn = user.role.name
    if rn in ("head","pm"):
        tasks = db.query(models.Task).all()
    elif rn == "employee":
        tasks = db.query(models.Task).filter(models.Task.assignee_id == user.id).all()
    else:
        return RedirectResponse(url="/projects")
    return templates.TemplateResponse("tasks.html", {"request": request, "user": user, "tasks": tasks})

@app.get("/tasks/new", response_class=HTMLResponse)
async def new_task_get(request: Request, project_id: int = None, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user or user.role.name not in ("head","pm"): return RedirectResponse(url="/tasks")
    sf = db.query(models.ProjectStatus).filter_by(name="Завершен").first()
    projects = db.query(models.Project).filter(models.Project.status_id != sf.id).all()
    employees = db.query(models.User).join(models.Role).filter(models.Role.name == "employee").all()
    return templates.TemplateResponse("task_form.html", {
        "request": request, "user": user, "projects": projects,
        "employees": employees, "preselect_project": project_id
    })

@app.post("/tasks/new")
async def new_task_post(
    request: Request,
    title: str = Form(...),
    description: str = Form(""),
    project_id: int = Form(...),
    assignee_id: int = Form(...),
    deadline: date = Form(...),
    cost: float = Form(0.0),
    db: Session = Depends(get_db)
):
    user = get_current_user(request, db)
    if not user or user.role.name not in ("head","pm"): return RedirectResponse(url="/login")
    st = db.query(models.TaskStatus).filter_by(name="К выполнению").first()
    task = models.Task(title=title, description=description, project_id=project_id,
                       assignee_id=assignee_id, deadline=deadline, cost=cost, status_id=st.id)
    db.add(task)
    db.commit()
    db.refresh(task)
    notify(db, assignee_id, f"Вам назначена задача: «{title}»", f"/projects/{project_id}")
    db.commit()
    ref = request.headers.get("referer", "/tasks")
    return RedirectResponse(url=ref, status_code=status.HTTP_302_FOUND)

@app.post("/tasks/{task_id}/complete")
async def complete_task(request: Request, task_id: int, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user: return RedirectResponse(url="/login")
    task = db.query(models.Task).filter(models.Task.id == task_id).first()
    sd = db.query(models.TaskStatus).filter_by(name="Выполнено").first()
    if task and (user.role.name in ("head","pm") or task.assignee_id == user.id):
        old_sid = task.status_id
        task.status_id = sd.id
        db.add(models.TaskHistory(task_id=task.id, old_status_id=old_sid, new_status_id=sd.id, changed_by=user.id))
        proj = task.project
        if proj and proj.pm_id:
            notify(db, proj.pm_id, f"Задача «{task.title}» выполнена исполнителем {user.full_name}", f"/projects/{proj.id}")
        db.commit()
    return RedirectResponse(url=request.headers.get("referer", "/tasks"), status_code=status.HTTP_302_FOUND)

# ─── Users / Workload ─────────────────────────────────────────────────────────

@app.get("/users", response_class=HTMLResponse)
async def list_users(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user or user.role.name not in ("head","pm"): return RedirectResponse(url="/projects")
    users = db.query(models.User).all()
    # Build workload: count active tasks per employee
    employees = db.query(models.User).join(models.Role).filter(models.Role.name == "employee").all()
    active_status_ids = [
        s.id for s in db.query(models.TaskStatus).filter(models.TaskStatus.name.in_(("К выполнению","В процессе"))).all()
    ]
    workload = {}
    for emp in employees:
        active = db.query(models.Task).filter(
            models.Task.assignee_id == emp.id,
            models.Task.status_id.in_(active_status_ids)
        ).count()
        total = db.query(models.Task).filter(models.Task.assignee_id == emp.id).count()
        workload[emp.id] = {"active": active, "total": total}
    return templates.TemplateResponse("users.html", {"request": request, "user": user, "users": users, "workload": workload})

# ─── Leads ────────────────────────────────────────────────────────────────────

@app.get("/leads", response_class=HTMLResponse)
async def list_leads(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user: return RedirectResponse(url="/login")
    rn = user.role.name
    if rn in ("head","pm"):
        leads = db.query(models.Lead).all()
    elif rn == "customer":
        leads = db.query(models.Lead).filter(models.Lead.customer_id == user.id).all()
    else:
        return RedirectResponse(url="/projects")
    doc_types = db.query(models.DocumentType).all()
    return templates.TemplateResponse("leads.html", {"request": request, "user": user, "leads": leads, "doc_types": doc_types})

@app.post("/leads/new")
async def create_lead(
    request: Request,
    title: str = Form(...),
    description: str = Form(""),
    budget: float = Form(0.0),
    desired_deadline: date = Form(None),
    tz_file: UploadFile = File(None),
    db: Session = Depends(get_db)
):
    user = get_current_user(request, db)
    if not user or user.role.name != "customer": return RedirectResponse(url="/login")
    tz_path = None
    if tz_file and tz_file.filename:
        ext = os.path.splitext(tz_file.filename)[1]
        fname = f"tz_{user.id}_{int(datetime.utcnow().timestamp())}{ext}"
        fpath = os.path.join(UPLOAD_DIR, fname)
        with open(fpath, "wb") as f:
            shutil.copyfileobj(tz_file.file, f)
        tz_path = fname
    lead = models.Lead(title=title, description=description, budget=budget,
                       desired_deadline=desired_deadline, customer_id=user.id, status="Новая", tz_file_path=tz_path)
    db.add(lead)
    db.commit()
    # Notify managers
    managers = db.query(models.User).join(models.Role).filter(models.Role.name.in_(("head","pm"))).all()
    for m in managers:
        notify(db, m.id, f"Новая заявка от {user.full_name}: «{title}»", "/leads")
    db.commit()
    return RedirectResponse(url="/leads", status_code=status.HTTP_302_FOUND)

@app.post("/leads/{lead_id}/approve")
async def approve_lead(request: Request, lead_id: int, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user or user.role.name not in ("head","pm"): return RedirectResponse(url="/login")
    lead = db.query(models.Lead).filter(models.Lead.id == lead_id).first()
    if lead and lead.status == "Новая":
        lead.status = "Одобрена"
        sn = db.query(models.ProjectStatus).filter_by(name="Новый").first()
        proj = models.Project(title=lead.title, description=lead.description, customer_id=lead.customer_id,
                              pm_id=user.id, start_date=date.today(), end_date=lead.desired_deadline,
                              planned_budget=lead.budget, status_id=sn.id)
        db.add(proj)
        db.commit()
        db.refresh(proj)
        notify(db, lead.customer_id, f"Ваша заявка «{lead.title}» одобрена — создан проект!", f"/projects/{proj.id}")
        db.commit()
    return RedirectResponse(url="/projects", status_code=status.HTTP_302_FOUND)

@app.post("/leads/{lead_id}/reject")
async def reject_lead(request: Request, lead_id: int, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user or user.role.name not in ("head","pm"): return RedirectResponse(url="/login")
    lead = db.query(models.Lead).filter(models.Lead.id == lead_id).first()
    if lead:
        lead.status = "Отклонена"
        notify(db, lead.customer_id, f"Заявка «{lead.title}» отклонена. Обратитесь к менеджеру.", "/leads")
        db.commit()
    return RedirectResponse(url="/leads", status_code=status.HTTP_302_FOUND)

# ─── Notifications ────────────────────────────────────────────────────────────

@app.get("/notifications", response_class=HTMLResponse)
async def notifications_page(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user: return RedirectResponse(url="/login")
    notifs = db.query(models.Notification).filter(models.Notification.user_id == user.id).order_by(models.Notification.created_at.desc()).limit(50).all()
    db.query(models.Notification).filter(models.Notification.user_id == user.id, models.Notification.is_read == False).update({"is_read": True})
    db.commit()
    return templates.TemplateResponse("notifications.html", {"request": request, "user": user, "notifs": notifs})

@app.get("/api/notifications/count")
async def notif_count(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user: return {"count": 0}
    cnt = db.query(models.Notification).filter(models.Notification.user_id == user.id, models.Notification.is_read == False).count()
    return {"count": cnt}

# ─── Export ───────────────────────────────────────────────────────────────────

@app.get("/projects/{project_id}/export/excel")
async def export_excel(request: Request, project_id: int, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user or user.role.name not in ("head","pm"): return RedirectResponse(url="/login")
    proj = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not proj: raise HTTPException(404)

    wb = Workbook()
    ws = wb.active
    ws.title = "Смета проекта"

    # Header styles
    h_fill = PatternFill("solid", fgColor="1E40AF")
    h_font = Font(name="Calibri", bold=True, color="FFFFFF", size=12)
    h_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # Project info
    ws.merge_cells("A1:F1")
    ws["A1"] = f"СМЕТА ПРОЕКТА: {proj.title}"
    ws["A1"].font = Font(name="Calibri", bold=True, size=14, color="1E293B")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    info_rows = [
        ("Статус:", proj.status.name),
        ("Дата начала:", str(proj.start_date)),
        ("Менеджер проекта:", proj.pm.full_name if proj.pm else "—"),
        ("Заказчик:", proj.customer.full_name if proj.customer else "—"),
        ("Плановый бюджет (₽):", f"{proj.planned_budget:,.0f}"),
    ]
    for i, (k, v) in enumerate(info_rows, start=2):
        ws[f"A{i}"] = k
        ws[f"A{i}"].font = Font(bold=True, name="Calibri", color="475569")
        ws[f"B{i}"] = v
        ws[f"B{i}"].font = Font(name="Calibri")

    # Tasks table header
    header_row = len(info_rows) + 3
    headers = ["№", "Задача", "Исполнитель", "Дедлайн", "Статус", "Стоимость (₽)"]
    col_widths = [5, 35, 25, 15, 15, 18]
    for col, (h, w) in enumerate(zip(headers, col_widths), start=1):
        cell = ws.cell(row=header_row, column=col, value=h)
        cell.fill = h_fill
        cell.font = h_font
        cell.alignment = h_align
        cell.border = border
        ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = w

    total_cost = 0
    for i, task in enumerate(proj.tasks, start=1):
        row = header_row + i
        row_data = [
            i, task.title,
            task.assignee.full_name if task.assignee else "—",
            str(task.deadline) if task.deadline else "—",
            task.status.name,
            task.cost
        ]
        bg = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        fill = PatternFill("solid", fgColor=bg)
        for col, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.fill = fill
            cell.border = border
            cell.font = Font(name="Calibri", size=11)
            if col == 1:
                cell.alignment = Alignment(horizontal="center")
            elif col == 6:
                cell.number_format = '#,##0.00'
        total_cost += task.cost

    # Total row
    total_row = header_row + len(proj.tasks) + 1
    ws.cell(row=total_row, column=5, value="ИТОГО:").font = Font(bold=True, name="Calibri")
    tc = ws.cell(row=total_row, column=6, value=total_cost)
    tc.font = Font(bold=True, name="Calibri", color="166534")
    tc.number_format = '#,##0.00'
    tc.fill = PatternFill("solid", fgColor="F0FDF4")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f"smeta_{proj.id}_{proj.title[:20].replace(' ','_')}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'}
    )

@app.get("/projects/{project_id}/export/word")
async def export_word(request: Request, project_id: int, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user or user.role.name not in ("head","pm"): return RedirectResponse(url="/login")
    proj = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not proj: raise HTTPException(404)

    doc = DocxDocument()

    # Title
    title_par = doc.add_heading("", 0)
    run = title_par.add_run(f"Техническое задание\n«{proj.title}»")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    doc.add_paragraph()

    # Info section
    info = doc.add_paragraph()
    info.add_run("Общие сведения о проекте\n").bold = True
    details = [
        ("Статус", proj.status.name),
        ("Дата начала", str(proj.start_date)),
        ("Дата окончания", str(proj.end_date) if proj.end_date else "В процессе"),
        ("Плановый бюджет", f"{proj.planned_budget:,.0f} ₽"),
        ("Менеджер проекта", proj.pm.full_name if proj.pm else "—"),
        ("Заказчик", proj.customer.full_name if proj.customer else "—"),
    ]
    for k, v in details:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{k}: ").bold = True
        p.add_run(v)

    doc.add_heading("Описание проекта", level=1)
    doc.add_paragraph(proj.description or "—")

    doc.add_heading("Перечень задач", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["№", "Задача", "Исполнитель", "Дедлайн", "Статус"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True

    for i, task in enumerate(proj.tasks, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = task.title
        row[2].text = task.assignee.full_name if task.assignee else "—"
        row[3].text = str(task.deadline) if task.deadline else "—"
        row[4].text = task.status.name

    doc.add_paragraph()
    doc.add_paragraph(f"Документ сформирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"tz_{proj.id}_{proj.title[:20].replace(' ','_')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'}
    )

# ─── 1C API ───────────────────────────────────────────────────────────────────

@app.get("/api/1c/sync-employees")
async def sync_employees_1c(db: Session = Depends(get_db)):
    users = db.query(models.User).join(models.Role).filter(models.Role.name == "employee").all()
    return {"status": "success", "data": [{"id": u.id, "name": u.full_name} for u in users]}
